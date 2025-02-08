#!/usr/bin/env python3
"""
combine.py
----------
本模块完成以下工作：
1. 遍历下载文件夹中 HLHAD 分析的结果文件，提取每个样本的 HLA 数据。
2. 结合 sample_info.xlsx 中的样本信息，根据 Company 和 Huben 匹配到对应记录，
   获取 Donor_ID（sample 列）和 LotNumber（lot 列）。
3. 为每个样本生成 PDF 报告，报告内容包含三个“表格”，格式如下：
   ┌─────────────────────────────┐
   │ LotNumber     Donor_ID      │   ← 表格1（2列，总宽450，每列225），标题行加粗、浅灰背景
   │ 12430         059K010       │
   ├─────────────────────────────┤
   │ HLA-A         HLA-B     HLA-C │   ← 表格2（3列，总宽450，每列150），标题行加粗、浅灰背景
   │ 02:01:01,30:01:01 07:02:01,13:02:01 06:02:01,07:02:01 │
   ├─────────────────────────────┤
   │ HLA-DQB1      HLA-DRB1   HLA-DPB1│  ← 表格3（3列，总宽450，每列150），标题行加粗、浅灰背景
   │ 02:01:01,30:01:01 07:02:01,13:02:01 06:02:01,07:02:01 │
   └─────────────────────────────┘
   这3个表格紧密相连，无间隔，看起来像一个整体。
4. 基于模板生成 Word 报告（内容部分不再赘述，可参照之前代码）。
5. 汇总所有样本数据生成一个 Excel 文件。
生成的 PDF 与 Word 文件均保存于 /home/huben/hlahd.1.7.0/sample/ 下，
Excel 文件命名为 {excel_base}.xlsx，其中 excel_base 取自下载文件夹名称中第二部分（例如 "20250125"）。
"""

import os
import glob
import pandas as pd
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
from reportlab.lib import colors
from docx import Document
from docx.shared import Pt

# ---------------------
# 参数设置
BASE_SAMPLE_DIR = "/home/huben/hlahd.1.7.0/sample"
SAMPLE_INFO_FILE = "/home/huben/hlahd.1.7.0/sample_info/sample_info.xlsx"
WORD_TEMPLATE = "/home/huben/hlahd.1.7.0/sample_info/HLA-typing.docx"
# 只关注的 HLA 基因（HLA-A, HLA-B, HLA-C, HLA-DQB1, HLA-DRB1, HLA-DPB1）
GENES = ["A", "B", "C", "DQB1", "DRB1", "DPB1"]

def find_download_folder(base_dir):
    """
    在 base_dir 下寻找下载文件夹（排除特定名称的目录）
    返回找到的第一个目录
    """
    for item in os.listdir(base_dir):
        full_path = os.path.join(base_dir, item)
        if os.path.isdir(full_path) and item not in ["result"] and not item.endswith(".pdf") and not item.endswith(".xlsx"):
            return full_path
    return None

def extract_hla_from_file(result_file_path):
    """
    解析最终结果文本文件，返回一个字典，键为基因（例如 "A"），值为等位基因串。
    若第二列为 "-"，则复制第一列。
    去除 "HLA-<gene>*" 前缀，仅保留等位基因号（如 "02:01:01"）。
    """
    hla_data = {}
    with open(result_file_path, "r") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            parts = line.split()
            if len(parts) < 3:
                continue
            gene = parts[0].strip()
            if gene not in GENES:
                continue
            allele1 = parts[1].strip()
            allele2 = parts[2].strip()
            if "*" in allele1:
                allele1 = allele1.split("*")[1]
            if "*" in allele2:
                allele2 = allele2.split("*")[1]
            if allele2 == "-":
                allele2 = allele1
            hla_data[gene] = f"{allele1},{allele2}"
    return hla_data

def generate_pdf(pdf_row, pdf_path):
    """
    生成 PDF 报告。报告内容由 3 个紧密相连的表格组成：
    
    表格1（2列，总宽450，每列225）：  
      - 第1行：标题 “LotNumber” 与 “Donor_ID”（加粗、浅灰背景）  
      - 第2行：对应数值  
    
    表格2（3列，总宽450，每列150）：  
      - 第1行：标题 “HLA-A”、“HLA-B”、“HLA-C”（加粗、浅灰背景）  
      - 第2行：对应等位基因  
    
    表格3（3列，总宽450，每列150）：  
      - 第1行：标题 “HLA-DQB1”、“HLA-DRB1”、“HLA-DPB1”（加粗、浅灰背景）  
      - 第2行：对应等位基因  
    
    三个表格依次添加，无间隔，看起来像一个整体。
    """
    doc = SimpleDocTemplate(pdf_path, pagesize=A4)
    elements = []
    
    # 表格1：2列，2行，宽度：[225,225]
    data1 = [
        ["LotNumber", "Donor_ID"],
        [pdf_row.get("LotNumber", ""), pdf_row.get("Donor_ID", "")]
    ]
    table1 = Table(data1, colWidths=[225, 225])
    style1 = TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
        ('TEXTCOLOR', (0,0), (-1,0), colors.black),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('FONTSIZE', (0,0), (-1,-1), 10),
        ('GRID', (0,0), (-1,-1), 1, colors.black)
    ])
    table1.setStyle(style1)
    elements.append(table1)
    
    # 表格2：3列，2行，宽度：[150,150,150]
    data2 = [
        ["HLA-A", "HLA-B", "HLA-C"],
        [pdf_row.get("A", ""), pdf_row.get("B", ""), pdf_row.get("C", "")]
    ]
    table2 = Table(data2, colWidths=[150, 150, 150])
    style2 = TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
        ('TEXTCOLOR', (0,0), (-1,0), colors.black),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('FONTSIZE', (0,0), (-1,-1), 10),
        ('GRID', (0,0), (-1,-1), 1, colors.black)
    ])
    table2.setStyle(style2)
    elements.append(table2)
    
    # 表格3：3列，2行，宽度：[150,150,150]
    data3 = [
        ["HLA-DQB1", "HLA-DRB1", "HLA-DPB1"],
        [pdf_row.get("DQB1", ""), pdf_row.get("DRB1", ""), pdf_row.get("DPB1", "")]
    ]
    table3 = Table(data3, colWidths=[150, 150, 150])
    style3 = TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
        ('TEXTCOLOR', (0,0), (-1,0), colors.black),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('FONTSIZE', (0,0), (-1,-1), 10),
        ('GRID', (0,0), (-1,-1), 1, colors.black)
    ])
    table3.setStyle(style3)
    elements.append(table3)
    
    # 三个表格紧密相连，不添加 Spacer
    doc.build(elements)

def format_date_from_folder(download_folder):
    """
    从下载文件夹名称（例如 "HBBIO-20250125-L-01-2025-01-281900"）中提取日期信息：
    取第二部分（如 "20250125"），然后取后6位（例如 "250125"），格式化为 "25年1月25日"
    """
    base_name = os.path.basename(download_folder)
    parts = base_name.split("-")
    if len(parts) < 2:
        return ""
    second = parts[1]  # e.g. "20250125"
    if len(second) < 6:
        return ""
    sub = second[-6:]  # e.g. "250125"
    year = sub[:2]
    month = str(int(sub[2:4]))
    day = str(int(sub[4:6]))
    formatted = f"{year}年{month}月{day}日"
    return formatted

def set_cell_background(cell, fill_color):
    """
    设置 Word 表格单元格背景颜色，使用 XML 操作方式。
    """
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), fill_color)))

def generate_word(pdf_row, template_path, word_output_path, formatted_date):
    """
    基于模板生成 Word 报告：
      - 加载模板文件 template_path
      - 遍历文档段落，若发现包含 "报告日期" 的段落，则将其内容替换为 "报告日期：<formatted_date>"（下划线）
      - 添加分页符，新页中添加标题 “三、结果分析”（加粗、字号14），
        然后依次添加3个表格：
         表格1：LotNumber 与 Donor_ID（2列）
         表格2：HLA-A, HLA-B, HLA-C（3列）
         表格3：HLA-DQB1, HLA-DRB1, HLA-DPB1（3列）
      - 保存 Word 文件至 word_output_path
    """
    doc = Document(template_path)
    
    # 修改包含“报告日期”的段落
    for para in doc.paragraphs:
        if "报告日期" in para.text:
            para.text = f"报告日期：{formatted_date}"
            for run in para.runs:
                run.font.underline = True
            break

    doc.add_page_break()
    
    # 添加标题“ 三、结果分析 ”
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("三、结果分析")
    title_run.bold = True
    title_run.font.size = Pt(14)
    doc.add_paragraph()  # 空行

    # 表格1：LotNumber 与 Donor_ID（2列2行，设置为宽450，即分别225）
    table1 = doc.add_table(rows=2, cols=2)
    table1.style = 'Table Grid'
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = "LotNumber"
    hdr_cells[1].text = "Donor_ID"
    for cell in hdr_cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
        set_cell_background(cell, "D3D3D3")
    data_cells = table1.rows[1].cells
    data_cells[0].text = pdf_row.get("LotNumber", "")
    data_cells[1].text = pdf_row.get("Donor_ID", "")
    
    # 表格2：HLA-A, HLA-B, HLA-C（3列2行，每列150）
    table2 = doc.add_table(rows=2, cols=3)
    table2.style = 'Table Grid'
    hdr_cells = table2.rows[0].cells
    hdr_cells[0].text = "HLA-A"
    hdr_cells[1].text = "HLA-B"
    hdr_cells[2].text = "HLA-C"
    for cell in hdr_cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
        set_cell_background(cell, "D3D3D3")
    data_cells = table2.rows[1].cells
    data_cells[0].text = pdf_row.get("A", "")
    data_cells[1].text = pdf_row.get("B", "")
    data_cells[2].text = pdf_row.get("C", "")
    
    # 表格3：HLA-DQB1, HLA-DRB1, HLA-DPB1（3列2行，每列150）
    table3 = doc.add_table(rows=2, cols=3)
    table3.style = 'Table Grid'
    hdr_cells = table3.rows[0].cells
    hdr_cells[0].text = "HLA-DQB1"
    hdr_cells[1].text = "HLA-DRB1"
    hdr_cells[2].text = "HLA-DPB1"
    for cell in hdr_cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
        set_cell_background(cell, "D3D3D3")
    data_cells = table3.rows[1].cells
    data_cells[0].text = pdf_row.get("DQB1", "")
    data_cells[1].text = pdf_row.get("DRB1", "")
    data_cells[2].text = pdf_row.get("DPB1", "")
    
    doc.save(word_output_path)

def main():
    # 1. 找到下载文件夹（例如 HBBIO-20250125-L-01-2025-01-281900）
    download_folder = find_download_folder(BASE_SAMPLE_DIR)
    if not download_folder:
        print("未在 {} 下找到下载文件夹！".format(BASE_SAMPLE_DIR))
        return
    print("下载文件夹：", download_folder)
    
    # 提取 Excel/Word 的基础名称：取下载文件夹名称中第二部分（例如 "20250125"）
    folder_parts = os.path.basename(download_folder).split("-")
    if len(folder_parts) < 2:
        print("下载文件夹名称格式异常！")
        return
    excel_base = folder_parts[1]
    excel_save_path = os.path.join(BASE_SAMPLE_DIR, excel_base + ".xlsx")
    
    # 计算报告日期
    formatted_date = format_date_from_folder(download_folder)
    
    # HLHAD 结果存放在 download_folder/result
    result_dir = os.path.join(download_folder, "result")
    if not os.path.isdir(result_dir):
        print("结果目录 {} 不存在！".format(result_dir))
        return
    
    # 读取 sample_info.xlsx
    try:
        sample_info_df = pd.read_excel(SAMPLE_INFO_FILE, engine='openpyxl')
    except Exception as e:
        print("读取 sample_info.xlsx 失败：", e)
        return
    
    pdf_data_rows = []  # 用于汇总 Excel 数据，每项为 dict
    sample_folders = [d for d in os.listdir(result_dir) if os.path.isdir(os.path.join(result_dir, d)) and d.startswith("JZ")]
    for sample_folder in sample_folders:
        sample_folder_full = os.path.join(result_dir, sample_folder)
        # 每个样本文件夹内部有一个 result 子目录，其中包含最终结果文件 *_final.result.txt
        inner_result_dir = os.path.join(sample_folder_full, "result")
        if not os.path.isdir(inner_result_dir):
            print("样本文件夹 {} 中未找到 result 子目录，跳过".format(sample_folder))
            continue
        txt_files = glob.glob(os.path.join(inner_result_dir, "*_final.result.txt"))
        if not txt_files:
            print("样本文件夹 {} 中未找到最终结果文件，跳过".format(sample_folder))
            continue
        final_result_file = txt_files[0]
        hla_dict = extract_hla_from_file(final_result_file)
        
        # 解析文件名以获取 Company 和样本标识
        # 例如：JZ25020604-009C250124-009C25012401_final.result.txt
        base_txt_name = os.path.basename(final_result_file)
        parts = base_txt_name.split("-")
        if len(parts) < 3:
            print("文件名格式异常：", base_txt_name)
            continue
        company = parts[1]  # 例如 "009C250124"
        third_part = parts[2]
        sample_id_full = third_part.split("_")[0]  # 例如 "009C25012401"
        huben_str = sample_id_full[-2:]
        try:
            huben_val = int(huben_str)
        except:
            print("无法转换 Huben 数值：", huben_str)
            continue
        
        # 在 sample_info.xlsx 中查找匹配记录：匹配 Company 和 Huben
        df_match = sample_info_df[(sample_info_df["Company"]==company) & (sample_info_df["Huben"]==huben_val)]
        if df_match.empty:
            print("未在 sample_info.xlsx 中找到 Company={} Huben={} 的记录，跳过".format(company, huben_val))
            continue
        record = df_match.iloc[0]
        donor_id = str(record["sample"]).strip()   # Donor_ID
        lot_number = str(record["lot"]).strip()      # LotNumber
        
        pdf_row = {
            "LotNumber": lot_number,
            "Donor_ID": donor_id,
            "A": hla_dict.get("A", ""),
            "B": hla_dict.get("B", ""),
            "C": hla_dict.get("C", ""),
            "DQB1": hla_dict.get("DQB1", ""),
            "DRB1": hla_dict.get("DRB1", ""),
            "DPB1": hla_dict.get("DPB1", "")
        }
        pdf_data_rows.append(pdf_row)
        
        # 生成 PDF 文件（保存至 BASE_SAMPLE_DIR，文件名为 Donor_ID.pdf）
        pdf_file_path = os.path.join(BASE_SAMPLE_DIR, donor_id + ".pdf")
        try:
            generate_pdf(pdf_row, pdf_file_path)
            print("生成 PDF：", pdf_file_path)
        except Exception as e:
            print("生成 PDF 失败（{}）：{}".format(pdf_file_path, e))
        
        # 生成 Word 文档（文件名为 {excel_base}_{Donor_ID}.docx）
        word_file_name = f"{excel_base}_{donor_id}.docx"
        word_file_path = os.path.join(BASE_SAMPLE_DIR, word_file_name)
        try:
            generate_word(pdf_row, WORD_TEMPLATE, word_file_path, formatted_date)
            print("生成 Word：", word_file_path)
        except Exception as e:
            print("生成 Word 失败（{}）：{}".format(word_file_path, e))
    
    # 生成汇总 Excel 文件
    if pdf_data_rows:
        df_summary = pd.DataFrame(pdf_data_rows, columns=["LotNumber", "Donor_ID", "A", "B", "C", "DQB1", "DRB1", "DPB1"])
        df_summary.rename(columns={"A": "HLA-A", "B": "HLA-B", "C": "HLA-C",
                                   "DQB1": "HLA-DQB1", "DRB1": "HLA-DRB1", "DPB1": "HLA-DPB1"}, inplace=True)
        try:
            df_summary.to_excel(excel_save_path, index=False, engine='openpyxl')
            print("生成 Excel 汇总文件：", excel_save_path)
        except Exception as e:
            print("生成 Excel 文件失败：", e)
    else:
        print("未收集到任何样本数据，未生成 Excel。")

if __name__ == "__main__":
    main()
